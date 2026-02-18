# Save this as setup_project.py
# Run it using: python setup_project.py
# (Requires: pip install python-docx)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = Document()
    
    # --- PROPOSAL SECTION ---
    doc.add_heading('FORMAL PROJECT PROPOSAL', 0)
    
    title = doc.add_heading('CuraAudit: Autonomous Multi-Agent Framework for Healthcare Financial Integrity', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Abstract
    doc.add_heading('1. Abstract', level=2)
    doc.add_paragraph(
        "CuraAudit is an autonomous auditing system designed for healthcare multi-tenant environments. "
        "The system utilizes a Multi-Agent System (MAS) architecture to bridge the gap between semantic policy "
        "documents (PDF contracts) and structured financial data (payment records). By deploying specialized "
        "AI agents for policy interpretation, data analysis, and conflict resolution, CuraAudit automates the "
        "detection of financial leakage, billing errors, and fraudulent activity in real-time."
    )

    # 2. Problem Statement
    doc.add_heading('2. Problem Statement', level=2)
    doc.add_paragraph(
        "The healthcare industry suffers from significant revenue loss due to administrative errors and "
        "unreconciled payment callbacks. Traditional rule-based systems are incapable of understanding "
        "complex contractual exceptions or identifying multi-step fraud patterns. Manual reconciliation is "
        "cost-prohibitive and slow, leading to financial instability and lack of transparency."
    )

    # 3. Project Objectives
    doc.add_heading('3. Project Objectives', level=2)
    obj = doc.add_paragraph(style='List Bullet')
    obj.add_run("Autonomous Reconciliation: Match external payment gateway logs with internal ledgers using agentic workflows.")
    obj = doc.add_paragraph(style='List Bullet')
    obj.add_run("Semantic Compliance: Implement RAG-based analysis of PDF care agreements to verify billing accuracy.")
    obj = doc.add_paragraph(style='List Bullet')
    obj.add_run("Explainable Auditing: Provide a human-readable 'Reasoning Chain' for every flagged discrepancy.")
    obj = doc.add_paragraph(style='List Bullet')
    obj.add_run("System Security: Implement robust Login/Signup and Role-Based Access Control (RBAC) to ensure data privacy.")

    # 4. Technical Architecture
    doc.add_heading('4. Technical Architecture', level=2)
    doc.add_paragraph(
        "Frontend: Next.js 14 (App Router) with Tailwind CSS.\n"
        "Backend: Node.js/TypeScript (GraphQL) and Python (FastAPI).\n"
        "AI Layer: LangGraph for agent orchestration, OpenAI/Gemini for reasoning.\n"
        "Data: MongoDB (Persistent storage) and ChromaDB (Vector store for RAG)."
    )

    # 5. Security & Multi-Tenancy
    doc.add_heading('5. Security & Multi-Tenancy', level=2)
    doc.add_paragraph(
        "The system includes a secure Authentication module (Login/Signup/JWT). "
        "Each facility is isolated via Multi-Tenant Scoping, ensuring Hospital A cannot access Hospital B’s records. "
        "AI agents are strictly scoped to the user’s organization context."
    )

    # --- TEAM TASK LIST SECTION ---
    doc.add_page_break()
    doc.add_heading('WEEK 1: TEAM TASK LIST', 0)
    
    # Task 1: You (The Backend Lead)
    doc.add_heading('Task 1: Project Lead (You)', level=2)
    doc.add_paragraph(
        "Goal: Setup the 'Brain' and Core Data.\n"
        "- Register and get API Keys (Google Gemini or Groq).\n"
        "- Prepare the 'Fake' Data: A MongoDB collection with 20 subscriptions and 20 payment records (some with intentional mistakes).\n"
        "- Setup the Python (FastAPI) folder structure where the AI agents will live."
    )

    # Task 2: Teammate 1 (Frontend Lead)
    doc.add_heading('Task 2: Frontend & UI Specialist', level=2)
    doc.add_paragraph(
        "Goal: Build the 'Home' for the project.\n"
        "- Initialize the Next.js 14 project.\n"
        "- Build the Login and Signup pages (use Shadcn UI).\n"
        "- Create the Basic Dashboard layout with a sidebar (Home, Audit Logs, AI Terminal, Settings)."
    )

    # Task 3: Teammate 2 (AI/Data Specialist)
    doc.add_heading('Task 3: AI & RAG Specialist', level=2)
    doc.add_paragraph(
        "Goal: Teach the AI how to read.\n"
        "- Gather 3 sample 'Hospital Care Agreements' as PDFs.\n"
        "- Setup a simple Python script to turn those PDFs into 'Embeddings' (text data the AI can search).\n"
        "- Create the 'System Prompts' for the first Agent (The Policy Clerk)."
    )

    # Task 4: Teammate 3 (DevOps/QA Specialist)
    doc.add_heading('Task 4: Quality & Documentation', level=2)
    doc.add_paragraph(
        "Goal: Keep the project organized.\n"
        "- Setup the GitHub Repository and invite all members.\n"
        "- Create the 'Project Readme' file.\n"
        "- Prepare the 'Friday Presentation Slides' template using the Proposal content."
    )

    doc.add_paragraph("\n[NOTE: Every member must have a working 'clickable' demo of their task by Thursday night.]")

    doc.save('CuraAudit_Project_Package.docx')
    print("Success! 'CuraAudit_Project_Package.docx' has been created.")

if __name__ == "__main__":
    create_proposal()